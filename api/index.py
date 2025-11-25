"""
Vercel serverless function for commodity news synthesis
"""
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
import os
import json
import httpx
from typing import List, Dict, Any
import datetime
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from io import BytesIO
import tempfile

# Import models (simplified version)
from pydantic import BaseModel
from typing import Optional, List

class NewsArticle(BaseModel):
    title: str
    content: str
    url: str
    published_date: Optional[str] = None
    source: Optional[str] = None

class ParaphraseRequest(BaseModel):
    articles: List[NewsArticle]

class ParaphraseResponse(BaseModel):
    synthesized_article: str
    title: str
    summary: str
    source_articles: List[NewsArticle]

# Initialize FastAPI app
app = FastAPI(title="AME Commodity News API")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize AI services
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyBUzBWciwGcb0m8DkYEMPkkwnqWD8ldE3g")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-1b5cc072ab944781a9c9ba0bf6936f79")

import google.generativeai as genai
genai.configure(api_key=GEMINI_API_KEY)

async def call_gemini_api(prompt: str) -> str:
    """Call Google Gemini API"""
    try:
        model = genai.GenerativeModel('gemini-2.0-flash-exp')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Gemini API error: {e}")
        return await call_deepseek_api(prompt)

async def call_deepseek_api(prompt: str) -> str:
    """Fallback to DeepSeek API"""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "deepseek-chat",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.7,
                    "max_tokens": 2000
                }
            )
            result = response.json()
            return result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"DeepSeek API error: {e}")
        raise HTTPException(status_code=500, detail="AI service unavailable")

def force_multiple_paragraphs(text: str) -> str:
    """Force text into multiple paragraphs with enhanced logic"""
    if not text or len(text.strip()) < 100:
        return text
    
    # Remove existing paragraph breaks
    text = text.replace('\n\n', ' ').replace('\n', ' ').strip()
    
    # If already has multiple clear paragraphs, return as-is
    paragraphs = [p.strip() for p in text.split('. ') if p.strip()]
    if len(paragraphs) < 3:
        # Split by character count to force 3 paragraphs
        char_per_para = len(text) // 3
        
        # Find natural break points near the split positions
        para1_end = char_per_para
        para2_end = char_per_para * 2
        
        # Adjust to sentence boundaries
        while para1_end < len(text) and text[para1_end] not in '.!?':
            para1_end += 1
        while para2_end < len(text) and text[para2_end] not in '.!?':
            para2_end += 1
            
        para1 = text[:para1_end + 1].strip()
        para2 = text[para1_end + 1:para2_end + 1].strip()
        para3 = text[para2_end + 1:].strip()
        
        return f"{para1}\n\n{para2}\n\n{para3}"
    else:
        # Join sentences into 3 paragraphs
        sentences_per_para = len(paragraphs) // 3
        para1 = '. '.join(paragraphs[:sentences_per_para]) + '.'
        para2 = '. '.join(paragraphs[sentences_per_para:sentences_per_para*2]) + '.'
        para3 = '. '.join(paragraphs[sentences_per_para*2:]) + '.'
        
        return f"{para1}\n\n{para2}\n\n{para3}"

@app.post("/api/news/paraphrase")
async def paraphrase_articles(request: ParaphraseRequest):
    """Synthesize news articles into professional commodity report"""
    try:
        if not request.articles:
            raise HTTPException(status_code=400, detail="No articles provided")
        
        # Build context from articles
        articles_content = []
        for i, article in enumerate(request.articles, 1):
            content = f"Article {i}: {article.title}\n{article.content[:2000]}..."
            articles_content.append(content)
        
        combined_content = "\n\n".join(articles_content)
        
        # Professional synthesis prompt
        synthesis_prompt = f"""
        As a senior commodities analyst and mining engineer with expertise in metals markets, synthesize the following news articles into a comprehensive professional report. Write in British English with a neutral, technical tone suitable for industry professionals.

        CRITICAL FORMATTING REQUIREMENT: Your response MUST contain exactly 3 distinct paragraphs separated by double line breaks. Each paragraph should be substantial (minimum 200 words).

        Focus on:
        - Technical analysis of production impacts, supply chain disruptions, or capacity changes
        - Quantitative data (production figures, price movements, tonnage, CAPEX commitments)
        - Market implications for commodity prices and supply dynamics
        - Regulatory or policy impacts on operations
        - Strategic implications for major market participants

        Maintain objectivity and avoid sensational language. Include specific figures, percentages, and technical details where available.

        Source articles:
        {combined_content}

        Write a comprehensive analysis in EXACTLY 3 paragraphs:
        """
        
        # Generate synthesis with multiple attempts for paragraph structure
        synthesized_content = await call_gemini_api(synthesis_prompt)
        
        # Force proper paragraph structure
        synthesized_content = force_multiple_paragraphs(synthesized_content)
        
        # Generate title
        title_prompt = f"Generate a professional, technical headline (maximum 80 characters) for this commodity market analysis:\n\n{synthesized_content[:500]}..."
        title = await call_gemini_api(title_prompt)
        title = title.strip().replace('"', '').replace('\n', ' ')[:80]
        
        # Generate summary
        summary_prompt = f"Create a concise 2-sentence executive summary of this commodity analysis:\n\n{synthesized_content[:800]}..."
        summary = await call_gemini_api(summary_prompt)
        summary = summary.strip().replace('\n', ' ')[:300]
        
        return ParaphraseResponse(
            synthesized_article=synthesized_content,
            title=title,
            summary=summary,
            source_articles=request.articles
        )
        
    except Exception as e:
        print(f"Error in paraphrase_articles: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/export/docx")
async def export_to_docx(request: ParaphraseResponse):
    """Export synthesized article to DOCX format"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading(request.title, 0)
        title.alignment = 1  # Center alignment
        
        # Add date
        doc.add_paragraph(f"Report Date: {datetime.datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph("")
        
        # Add summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(request.summary)
        doc.add_paragraph("")
        
        # Add main content
        doc.add_heading("Market Analysis", level=1)
        paragraphs = request.synthesized_article.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Add references section
        doc.add_page_break()
        doc.add_heading("Source References", level=1)
        for i, article in enumerate(request.source_articles, 1):
            ref_para = doc.add_paragraph()
            ref_para.add_run(f"{i}. ").bold = True
            ref_para.add_run(f"{article.title}\n")
            ref_para.add_run(f"Source: {article.source or 'Unknown'}\n")
            ref_para.add_run(f"URL: {article.url}\n")
            if article.published_date:
                ref_para.add_run(f"Published: {article.published_date}\n")
            ref_para.add_run("\n")
        
        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Generate filename
        safe_title = "".join(c for c in request.title if c.isalnum() or c in (' ', '_')).rstrip()
        filename = f"AME_Report_{safe_title[:30]}_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        
        return Response(
            content=docx_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        print(f"Error in export_to_docx: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/export/pdf")
async def export_to_pdf(request: ParaphraseResponse):
    """Export synthesized article to PDF format"""
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1
        )
        
        # Add title
        story.append(Paragraph(request.title, title_style))
        story.append(Spacer(1, 12))
        
        # Add date
        story.append(Paragraph(f"Report Date: {datetime.datetime.now().strftime('%d %B %Y')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Add summary
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        story.append(Paragraph(request.summary, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Add main content
        story.append(Paragraph("Market Analysis", styles['Heading2']))
        paragraphs = request.synthesized_article.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 12))
        
        story.append(Spacer(1, 30))
        
        # Add references
        story.append(Paragraph("Source References", styles['Heading2']))
        for i, article in enumerate(request.source_articles, 1):
            ref_text = f"<b>{i}.</b> {article.title}<br/>"
            ref_text += f"Source: {article.source or 'Unknown'}<br/>"
            ref_text += f"URL: {article.url}<br/>"
            if article.published_date:
                ref_text += f"Published: {article.published_date}<br/>"
            ref_text += "<br/>"
            story.append(Paragraph(ref_text, styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        
        # Generate filename
        safe_title = "".join(c for c in request.title if c.isalnum() or c in (' ', '_')).rstrip()
        filename = f"AME_Report_{safe_title[:30]}_{datetime.datetime.now().strftime('%Y%m%d')}.pdf"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        print(f"Error in export_to_pdf: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api")
async def root():
    """API health check"""
    return {"message": "AME Commodity News API", "status": "active"}

# For Vercel
handler = app